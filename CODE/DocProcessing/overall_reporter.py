"""
overall_reporter_anthropic.py
· 使用 Anthropic API（Claude 3 系列）生成合规总体报告
· 输出:
    ├── <doc>_overall.json   机器可读结构
    └── <doc>_overall.docx   格式化 Word
依赖:
    pip install anthropic python-docx

重要说明:
    使用RequirementID (1-35) 作为主键追踪风险类别，避免因文本名称差异导致的匹配问题
"""

from __future__ import annotations
import os, sys, json, textwrap, collections, re
import numpy as np
from pathlib import Path
from typing import Dict, List
from datetime import datetime
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from prompt import REGULATORY_FRAMEWORK

# 为导入可视化工具添加路径
BASE_DIR = Path(__file__).resolve().parents[1]
if str(BASE_DIR) not in sys.path:
    sys.path.append(str(BASE_DIR))


# ───────────────────────── Anthropic 调用辅助 ─────────────────────────
def _call_anthropic(system_msg: str, user_msg: str,
                    model="claude-opus-4-20250514",
                    temperature=0.2,
                    max_tokens=1024,
                    api_key: str | None = None) -> str:
    """
    Wrapper: 返回纯字符串（去掉 ```json``` 包裹）
    """
    api_key='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'
    client = anthropic.Anthropic(api_key=api_key)
    
    # 判断是否需要JSON格式
    needs_json = "JSON" in user_msg or "json" in user_msg
    
    if needs_json:
        enhanced_user_msg = (
            user_msg +
            "\n\n请确保返回有效的JSON格式，不要包含markdown代码块标记，也不要额外解释。"
        )
    else:
        enhanced_user_msg = user_msg
        
    msg = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        system=system_msg,
        messages=[{"role": "user", "content": enhanced_user_msg}],
    )
    # content 可能是 list(blocks)
    if isinstance(msg.content, list):
        content = "".join(b.text for b in msg.content if hasattr(b, "text"))
    else:
        content = str(msg.content)
    content = content.strip()
    
    # 只有在需要JSON时才清理JSON标记
    if needs_json:
        if content.startswith("```json") and content.endswith("```"):
            content = content[7:-3].strip()
        elif content.startswith("```") and content.endswith("```"):
            content = content[3:-3].strip()
    
    return content


# ───────────────────────── JSON 解析辅助 ─────────────────────────
def _safe_json_loads(text: str) -> Dict:
    """更稳健地解析可能被额外文本包裹的JSON字符串"""
    text = text.strip()
    
    # 尝试直接解析
    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        # 记录原始错误
        original_error = str(e)
        
    # 尝试提取JSON部分
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        candidate = text[start:end + 1]
        
        # 尝试修复常见的JSON格式问题
        try:
            # 1. 修复缺失的逗号（在 "}" 或 "]" 后面跟着 " 的情况）
            candidate = re.sub(r'([}\]])\s*"', r'\1, "', candidate)
            
            # 2. 修复数组中缺失的逗号
            candidate = re.sub(r'"\s+"', '", "', candidate)
            
            # 3. 修复对象中缺失的逗号（在引号后跟着新的键）
            candidate = re.sub(r'"\s*\n\s*"', '",\n"', candidate)
            
            # 4. 移除末尾多余的逗号
            candidate = re.sub(r',\s*}', '}', candidate)
            candidate = re.sub(r',\s*]', ']', candidate)
            
            # 5. 尝试解析修复后的JSON
            return json.loads(candidate)
        except json.JSONDecodeError:
            # 如果还是失败，尝试更激进的修复
            try:
                # 替换单引号为双引号
                candidate = re.sub(r"'([^']*)':", r'"\1":', candidate)
                candidate = re.sub(r":\s*'([^']*)'", r': "\1"', candidate)
                
                # 确保布尔值小写
                candidate = candidate.replace('True', 'true').replace('False', 'false')
                
                # 处理None值
                candidate = candidate.replace('None', 'null')
                
                return json.loads(candidate)
            except json.JSONDecodeError:
                pass
    
    # 如果所有尝试都失败，抛出原始错误
    raise json.JSONDecodeError(
        f"Failed to parse JSON. Original error: {original_error}",
        text, 0
    )


# ───────────────────────── 分析报告解析 ─────────────────────────
def _parse_analysis_report(text: str) -> tuple[list[list[str]], list[list[str]]]:
    """将文本形式的分析报告解析为表格数据"""
    import re

    llm_rows: list[list[str]] = []
    cat_rows: list[list[str]] = []


    llm_re = re.compile(
        r"^([A-Z]+):\s*$\n\s+- 平均得分: ([0-9.]+)\s*$\n\s+- 覆盖率: ([0-9.]+)%\s*$\n\s+- 高分项目数 .*: (\d+)",
        re.MULTILINE,
    )
    for m in llm_re.finditer(text):
        llm_rows.append([m.group(1), m.group(2), m.group(3) + "%", m.group(4)])


    cat_block_re = re.compile(
        r"^([一二三四五六七八]、[^:]+):\n((?:\s*[a-z]+: [^\n]+\n)+)",
        re.MULTILINE,
    )
    provider_re = re.compile(
        r"\s*(deepseek|openai|anthropic):\s*平均([0-9.]+)分,\s*最高([0-9.]+)分,\s*覆盖([0-9.]+)%",
    )

    for block in cat_block_re.finditer(text):
        cat = block.group(1)
        lines = block.group(2)
        for pm in provider_re.finditer(lines):
            cat_rows.append([cat, pm.group(1), pm.group(2), pm.group(3), pm.group(4) + "%"])

    return llm_rows, cat_rows


def _insert_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """在文档中插入带标题的表格"""

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E0E0E0')
        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return table


# ───────────────────────── 创建ID映射 ─────────────────────────
def _create_id_mappings():
    """
    创建ID到名称的映射和名称到ID的模糊匹配映射
    
    使用RequirementID (1-35) 作为主键的原因：
    - 避免因不同LLM返回的名称差异（如"与"和"和"的区别）导致的匹配失败
    - 确保所有36个风险子类别都被准确追踪和分析
    - 支持通过ID或名称查找要求
    """
    id_to_name = {}
    name_to_id = {}
    
    for cat, items in REGULATORY_FRAMEWORK.items():
        for item in items:
            id_to_name[item["number"]] = {
                "category": cat,
                "name": item["name"],
                "scope": item["scope"],
                "keyPoints": item["keyPoints"]
            }
            # 创建多个可能的名称变体用于匹配
            name_to_id[item["name"]] = item["number"]
            # 简化版本（去除标点）
            simplified = item["name"].replace("（", "(").replace("）", ")").replace("、", "")
            name_to_id[simplified] = item["number"]
    
    return id_to_name, name_to_id

def _find_requirement_id(name: str, name_to_id: dict) -> int | None:
    """根据名称查找requirement ID，支持模糊匹配"""
    # 精确匹配
    if name in name_to_id:
        return name_to_id[name]
    
    # 模糊匹配 - 查找包含关系
    for stored_name, req_id in name_to_id.items():
        if name in stored_name or stored_name in name:
            return req_id
    
    # 如果还找不到，尝试提取数字
    import re
    numbers = re.findall(r'\d+', name)
    if numbers and 1 <= int(numbers[0]) <= 35:
        return int(numbers[0])
    
    return None

# ───────────────────────── 数据聚合 ─────────────────────────
_COV_RANK = {
    "未提及": 0,
    "不适用": 0,
    "未覆盖": 1,
    "部分覆盖": 2,
    "完全覆盖": 3,
}

def _gather(json_path: Path):
    """解析综合 JSON → 覆盖矩阵、发现、建议"""
    raw = json.loads(Path(json_path).read_text(encoding="utf-8"))
    
    # 创建映射
    id_to_name, name_to_id = _create_id_mappings()
    
    # 使用ID作为key的覆盖矩阵
    cov_by_id = {}  # {req_id: {"coverage": lvl, "category": cat, "name": name}}
    findings, advice = [], []
    
    # 使用ID组织的详细数据
    detailed_data_by_id = collections.defaultdict(list)

    for pdata in raw["LLM分析结果"].values():
        # 覆盖
        for cat, items in pdata["详细分析"].items():
            for it in items:
                # 尝试获取requirement ID
                req_id = it.get("框架要求编号")
                if not req_id:
                    # 如果没有ID，尝试通过名称查找
                    req_id = _find_requirement_id(it.get("框架要求名称", ""), name_to_id)
                
                if req_id and req_id in id_to_name:
                    lvl = it["法规覆盖情况"]
                    
                    # 更新覆盖情况
                    if (
                        req_id not in cov_by_id
                        or _COV_RANK.get(lvl, 0)
                        > _COV_RANK.get(cov_by_id[req_id]["coverage"], 0)
                    ):
                        cov_by_id[req_id] = {
                            "coverage": lvl,
                            "category": id_to_name[req_id]["category"],
                            "name": id_to_name[req_id]["name"]
                        }
                    
                    # 收集详细数据
                    detailed_data_by_id[req_id].append({
                        "覆盖情况": lvl,
                        "法规要求内容": it.get("法规要求内容", []),
                        "实施要求": it.get("实施要求", ""),
                        "处罚措施": it.get("处罚措施", "")
                    })
                
        findings.extend(pdata.get("关键发现", []))
        advice.extend(pdata.get("合规建议", []))

    # 转换回按类别组织的格式，确保所有36个子类别都被包含
    cov = collections.defaultdict(dict)
    detailed_data = collections.defaultdict(lambda: collections.defaultdict(list))
    
    # 确保所有36个要求都被包含
    for req_id, info in id_to_name.items():
        cat = info["category"]
        name = info["name"]
        
        if req_id in cov_by_id:
            cov[cat][name] = cov_by_id[req_id]["coverage"]
            detailed_data[cat][name] = detailed_data_by_id[req_id]
        else:
            # 未提及的要求标记为"未覆盖"
            cov[cat][name] = "未覆盖"
            detailed_data[cat][name] = [{
                "覆盖情况": "未覆盖",
                "法规要求内容": [],
                "实施要求": "该法规未涉及此要求",
                "处罚措施": "不适用"
            }]

    findings = list(dict.fromkeys(findings))  # 去重保持顺序
    advice   = list(dict.fromkeys(advice))
    return cov, findings, advice, detailed_data


# ───────────────────────── 逐大类评估 ─────────────────────────
def _build_category_reports(cov, findings, advice, detailed_data,
                            model="claude-opus-4-20250514") -> List[Dict]:
    reports = []
    
    # 构建子类别详细信息字符串
    def build_subcategory_details(cat, sub_map):
        details = []
        # 检查detailed_data中是否有该类别
        if cat in detailed_data:
            for sub_name, data_list in detailed_data[cat].items():
                sub_info = f"\n子类别：{sub_name}\n"
                for item in data_list:
                    if item["法规要求内容"]:
                        sub_info += f"  法规要求：\n"
                        for req in item["法规要求内容"]:
                            sub_info += f"    - 条款{req.get('条款编号', '')}：{req.get('具体要求', '')}\n"
                    if item["实施要求"]:
                        sub_info += f"  实施要求：{item['实施要求']}\n"
                    if item["处罚措施"]:
                        sub_info += f"  处罚措施：{item['处罚措施']}\n"
                details.append(sub_info)
        
        # 如果没有详细数据，返回说明
        if not details:
            return "该法规对此大类没有具体要求。"
        
        return "".join(details)
    
    # 确保所有8个大类都被分析
    for cat_name in REGULATORY_FRAMEWORK.keys():
        # 如果某个大类在cov中不存在，创建空的覆盖信息
        if cat_name not in cov:
            cov[cat_name] = {}
        
        sub_map = cov[cat_name]
        # 获取该大类的子类别信息
        cat_subcategories = REGULATORY_FRAMEWORK.get(cat_name, [])
        subcategory_info = "\n".join([
            f"- {sub['number']}. {sub['name']} (关注点: {sub['keyPoints']})"
            for sub in cat_subcategories
        ])
        
        # 确保所有子类别都在sub_map中（即使是"未覆盖"）
        for sub in cat_subcategories:
            if sub['name'] not in sub_map:
                sub_map[sub['name']] = "未覆盖"
        
        # 构建详细的子类别数据
        detailed_info = build_subcategory_details(cat_name, sub_map)
        
        prompt = textwrap.dedent(f"""
        你是一名资深法律分析专家，专注于企业境外投资合规法律分析。

        **重要说明**：
        1. 这是一个法律分析任务，你需要分析该法规对企业在特定风险类别下的要求，而不是评判法规本身的质量或完整性。
        2. 你的任务是理解和解释法规对企业的具体要求，如果法规对某个风险子类别没有要求，请明确说明"该法规对此子类别无相关要求"。
        3. 多个来源的信息可能重复、交叉或矛盾，请基于你的专业知识整合这些信息，识别真实有效的要求。
        4. 我们使用RequirementID (1-35)作为主键追踪各个风险子类别，确保分析的准确性。

        **当前分析的风险大类**：{cat_name}
        
        **该大类包含的子类别**：
        {subcategory_info}

        **收集到的覆盖情况**：
        {json.dumps(sub_map, ensure_ascii=False)}

        **该大类的详细法规要求**：
        {detailed_info}

        **整体关键发现**：{findings}
        **整体合规建议**：{advice}

        请返回以下格式的JSON（注意字段必须完全匹配）:
        {{
          "Category": "{cat_name}",
          "CategoryLawAnalysis": "针对该大类，综合分析法规的整体要求和规范重点（300-400字）",
          "SubCategoryAnalysis": {{
              "子类名称": {{
                  "Coverage": "覆盖等级（未覆盖/部分覆盖/完全覆盖）",
                  "LawRequirements": "法规对该子类的具体要求说明（如无要求，明确说明'该法规对此子类别无相关要求'）（200-300字）",
                  "KeyProvisions": ["关键条款1", "关键条款2"],
                  "policy_document_needed": ["管理制度1", "管理制度2"]
                  "management_system_needed": ["管理体系1", "管理体系2"],
                  "Prohibited_List": ["禁止事项1", "禁止事项2"]
                  "Report_needed": ["报告1", "报告2"],
                  "data_info_needed": ["信息1", "信息2"],
                  "CompliancePoints": ["合规要点1", "合规要点2"]
              }},
              ...（所有子类都需要分析）
          }},
          "CategoryComplianceGuidance": "基于法规要求，企业在该大类下应采取的合规措施建议（150-200字）"
        }}
        
        注意：
        - 对于每个子类别，即使法规没有相关要求，也要在SubCategoryAnalysis中包含该子类别
        - 明确区分"法规要求"（法规规定企业必须做什么）和"合规建议"（为满足法规要求，企业应该如何做）
        - 请确保JSON格式正确，特别注意逗号的使用
        """)
        
        # 尝试多次获取正确的JSON
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                raw = _call_anthropic(
                    system_msg="你是专业的法律分析专家，必须返回有效的JSON格式，严格按照要求的字段结构。",
                    user_msg=prompt,
                    model=model,
                    max_tokens=6000,  # 增加token限制
                )
                parsed_report = _safe_json_loads(raw)
                reports.append(parsed_report)
                break
            except json.JSONDecodeError as e:
                print(f"尝试 {attempt + 1}/{max_attempts} - 解析类别 {cat_name} 的JSON响应时出错：{str(e)}")
                if attempt == max_attempts - 1:
                    # 最后一次尝试失败，创建默认报告
                    default_report = {
                        "Category": cat_name,
                        "CategoryLawAnalysis": f"由于技术原因，无法生成{cat_name}的详细分析。",
                        "SubCategoryAnalysis": {},
                        "CategoryComplianceGuidance": "建议重新运行分析以获取完整的合规指导。"
                    }
                    for sub in cat_subcategories:
                        default_report["SubCategoryAnalysis"][sub['name']] = {
                            "Coverage": sub_map.get(sub['name'], "未覆盖"),
                            "LawRequirements": "分析数据暂时不可用",
                            "KeyProvisions": [],
                            "CompliancePoints": []
                        }
                    reports.append(default_report)
                    
    return reports


# ───────────────────────── Word 导出 ─────────────────────────
def _export_word(report: Dict, out_file: Path, image_dir: Path | None = None, json_path: Path | None = None):
    """生成格式化的Word文档，包含适当的中文字体和表格样式"""

    doc = Document()
    
    # 设置文档默认字体
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    
    # 设置文档默认中文字体
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 创建标题样式
    def create_heading_style(level, size):
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.size = Pt(size)
        style.font.bold = True
        if level == 1:
            style.font.color.rgb = RGBColor(0, 0, 0)
        else:
            style.font.color.rgb = RGBColor(51, 51, 51)
    
    create_heading_style(1, 16)
    create_heading_style(2, 14)
    create_heading_style(3, 12)
    
    # 添加文档标题
    title = doc.add_heading(f"{report['DocumentTitle']} \n 法规要求分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = 'Arial'
    title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.runs[0].font.size = Pt(20)
    
    
    # 添加报告说明
    doc.add_heading("报告说明", level=1)
    para = doc.add_paragraph(
        "本报告旨在分析该法规对企业境外投资在8大风险类别和36个子类别下的具体要求。"
        "报告重点说明法规规定了什么、要求企业做什么，以及企业应如何满足这些要求。"
    )
    para.paragraph_format.first_line_indent = Inches(0.5)
    para.paragraph_format.space_after = Pt(12)

    # 法规整体分析与合规实施建议（合并为一章）
    doc.add_heading("法规整体分析与合规实施建议", level=1)
    
    # 添加方法论说明
    doc.add_heading("评分方法与分析框架说明", level=2)
    
    # 方法论介绍段落
    methodology_intro = doc.add_paragraph(
        "本报告采用综合评分方法，对法规在企业境外投资风险管理各个维度的覆盖程度进行量化分析。"
        "评分体系旨在帮助企业识别法规要求的重点领域、评估合规管理的优先级，并为建立全面的境外投资风险管理体系提供决策依据。"
    )
    methodology_intro.paragraph_format.first_line_indent = Inches(0.5)
    methodology_intro.paragraph_format.space_after = Pt(12)
    
    # 评分标准说明
    doc.add_heading("评分标准", level=3)
    scoring_criteria = doc.add_paragraph()
    scoring_criteria.add_run("评分采用0-100分制，根据法规对各项风险管理要求的覆盖程度进行评定：\n")
    scoring_criteria.paragraph_format.first_line_indent = Inches(0.5)
    
    # 创建评分标准表格
    score_table = doc.add_table(rows=1, cols=3)
    score_table.style = 'Light List Accent 1'
    
    # 表头
    hdr_cells = score_table.rows[0].cells
    headers = ["覆盖等级", "分值范围", "含义说明"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 评分标准数据
    score_standards = [
        ("完全覆盖", "100分", "法规对该风险管理要求有明确、详细的规定，并提供了具体的实施指引"),
        ("部分覆盖", "75分", "法规对该风险管理要求有相关规定，但规定较为原则性或覆盖不够全面"),
        ("低度覆盖", "15分", "法规仅对该风险管理要求有少量提及或间接涉及"),
        ("未覆盖", "0分", "法规未涉及该风险管理要求")
    ]
    
    for level, score, desc in score_standards:
        row_cells = score_table.add_row().cells
        row_cells[0].text = level
        row_cells[1].text = score
        row_cells[2].text = desc
        
        # 设置对齐
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 添加空行
    
    # 分析框架说明
    doc.add_heading("风险管理框架", level=3)
    framework_intro = doc.add_paragraph(
        "本报告基于企业境外投资风险管理最佳实践，构建了8大类、35个子类的全面风险管理框架，"
        "涵盖了企业境外投资全生命周期的主要风险领域："
    )
    framework_intro.paragraph_format.first_line_indent = Inches(0.5)
    framework_intro.paragraph_format.space_after = Pt(8)
    
    # 8大类说明
    categories_list = doc.add_paragraph()
    category_descriptions = [
        "**一、治理与战略**：涵盖海外业务治理结构、董事会职责、战略规划等",
        "**二、全面风险管理**：包括风险管理体系、风险评估、监测预警等",
        "**三、合规与法律**：涉及合规管理体系、反腐败、制裁、数据保护等",
        "**四、财务与市场风险**：包含汇率、商品价格、资金流动性管理等",
        "**五、运营与HSE**：覆盖环境保护、生产安全、供应链管理等",
        "**六、安全与危机**：涉及安全防护、应急管理、业务连续性等",
        "**七、信息与网络安全**：包括网络安全、数据安全、技术系统保护等",
        "**八、社会责任与人力**：涵盖社会责任、文化融合、人力资源管理等"
    ]
    
    for i, desc in enumerate(category_descriptions, 1):
        para = doc.add_paragraph()
        # 分离加粗部分和普通部分
        parts = desc.split("：", 1)
        bold_part = parts[0].replace("**", "")
        normal_part = "：" + parts[1] if len(parts) > 1 else ""
        
        # 添加加粗部分
        run = para.add_run(bold_part)
        run.font.bold = True
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加普通部分
        run = para.add_run(normal_part)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()  # 添加空行
    
    # 分析目的说明
    purpose_para = doc.add_paragraph(
        "通过对法规在上述框架下的覆盖程度进行系统评分，企业可以：\n"
        "1. 识别法规重点监管的风险领域\n"
        "2. 发现法规覆盖的空白点，需要企业自主加强管理\n"
        "3. 制定有针对性的合规管理措施\n"
        "4. 优化资源配置，确保重点领域的合规投入"
    )
    purpose_para.paragraph_format.first_line_indent = Inches(0.5)
    purpose_para.paragraph_format.space_after = Pt(12)
    
    # 添加说明文字
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        "以下通过热力图和数据表格展示法规在各风险管理维度的覆盖情况分析结果："
    )
    note_run.font.italic = True
    note_run.font.size = Pt(11)
    note_para.paragraph_format.space_after = Pt(12)
    
    # 添加分隔线
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()

    # 添加综合分析表格
    # 先计算类别统计和要求排名
    
    # 收集分数数据
    if json_path and json_path.exists():
        json_data = json.loads(json_path.read_text(encoding="utf-8"))
    else:
        # 如果没有json_path，尝试从image_dir找到对应的综合分析结果文件
        json_data = None
        if image_dir:
            possible_json = image_dir / f"{report['DocumentTitle']}_综合分析结果.json"
            if possible_json.exists():
                json_data = json.loads(possible_json.read_text(encoding="utf-8"))
    
    if json_data:
        id_to_name, _ = _create_id_mappings()
        
        scores_by_llm = {}
        scores_by_category = collections.defaultdict(lambda: collections.defaultdict(list))
        
        for llm_name, llm_data in json_data["LLM分析结果"].items():
            scores_by_llm[llm_name] = {}
            
            for cat_name, items in llm_data["详细分析"].items():
                for item in items:
                    req_id = item.get("框架要求编号")
                    if req_id and req_id in id_to_name:
                        coverage = item["法规覆盖情况"]
                        if coverage == "完全覆盖":
                            score = 100
                        elif coverage == "部分覆盖":
                            score = 75
                        elif coverage == "未覆盖":
                            score = 15
                        else:
                            score = 0
                        
                        scores_by_llm[llm_name][req_id] = score
                        scores_by_category[cat_name][llm_name].append(score)
        
        # 计算类别统计
        category_stats = {}
        for cat_name in REGULATORY_FRAMEWORK.keys():
            all_scores = []
            for llm_name in scores_by_category[cat_name]:
                all_scores.extend(scores_by_category[cat_name][llm_name])
            
            if all_scores:
                category_stats[cat_name] = {
                    "avg": np.mean(all_scores),
                    "max": max(all_scores),
                    "min": min(all_scores)
                }
            else:
                category_stats[cat_name] = {"avg": 0, "max": 0, "min": 0}
        
        # 添加类别分析表格
        doc.add_heading("风险类别覆盖度分析", level=2)
        cat_table = doc.add_table(rows=1, cols=4)
        cat_table.style = 'Light Grid Accent 1'
        
        # 设置表头
        headers = ["风险类别", "平均分", "最高分", "最低分"]
        hdr_cells = cat_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E0E0E0')
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # 添加数据
        for cat_name, stats in category_stats.items():
            row_cells = cat_table.add_row().cells
            row_cells[0].text = cat_name
            row_cells[1].text = f"{stats['avg']:.1f}"
            row_cells[2].text = f"{stats['max']:.0f}"
            row_cells[3].text = f"{stats['min']:.0f}"
            
            # 设置对齐
            for i in range(1, 4):
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 添加空行
        
        # 计算每个要求的平均分并排序
        requirement_scores = []
        for req_id, info in id_to_name.items():
            scores = []
            for llm_name in scores_by_llm:
                if req_id in scores_by_llm[llm_name]:
                    scores.append(scores_by_llm[llm_name][req_id])
                else:
                    scores.append(0)
            
            avg_score = sum(scores) / len(scores) if scores else 0
            requirement_scores.append({
                "id": req_id,
                "name": f"{req_id}. {info['name']}",
                "category": info['category'],
                "avg_score": avg_score
            })
        
        requirement_scores.sort(key=lambda x: x["avg_score"], reverse=True)
        
        # 添加详细要求覆盖度表格
        doc.add_heading("详细要求覆盖度排名", level=2)
        req_table = doc.add_table(rows=1, cols=4)
        req_table.style = 'Light Grid Accent 1'
        
        # 设置表头
        headers = ["排名", "要求名称", "所属类别", "平均分"]
        hdr_cells = req_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E0E0E0')
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # 添加所有要求的数据
        for rank, req in enumerate(requirement_scores, 1):
            row_cells = req_table.add_row().cells
            row_cells[0].text = str(rank)
            row_cells[1].text = req['name']
            row_cells[2].text = req['category']
            row_cells[3].text = f"{req['avg_score']:.1f}"
            
            # 设置对齐和颜色
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 根据分数设置颜色
            score_run = row_cells[3].paragraphs[0].runs[0]
            score_run.font.bold = True
            if req['avg_score'] >= 75:
                score_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif req['avg_score'] >= 25:
                score_run.font.color.rgb = RGBColor(255, 140, 0)  # 橙色
            else:
                score_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        
        # 设置列宽
        for i, width in enumerate([0.8, 4.0, 2.0, 1.0]):
            for cell in req_table.columns[i].cells:
                cell.width = Inches(width)
        
        doc.add_paragraph()  # 添加空行

    # 如有热力图和分析报告，插入于正文之前
    if image_dir:
        cat_img = next(Path(image_dir).glob("*分类汇总热力图.png"), None)
        det_img = next(Path(image_dir).glob("*详细热力图.png"), None)
        desc_map = {
            cat_img: "图：8大类风险法规防控要求覆盖热力图（分值越高要求越严格）",
            det_img: "图：36小类风险法规防控要求覆盖热力图（分值越高要求越严格）",
        }
        for img in [cat_img, det_img]:
            if img and img.exists():
                doc.add_picture(str(img), width=Inches(6))
                p = doc.add_paragraph(desc_map[img])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True

        txt_file = next(Path(image_dir).glob("*分析报告.txt"), None)
        if txt_file and txt_file.exists():
            llm_rows, cat_rows = _parse_analysis_report(txt_file.read_text(encoding="utf-8"))
            if llm_rows:
                doc.add_heading("得分汇总", level=2)
                _insert_table(doc, ["LLM", "平均得分", "覆盖率", "高分项目数"], llm_rows)
            if cat_rows:
                doc.add_heading("各类别得分概览", level=2)
                _insert_table(doc, ["类别", "LLM", "平均得分", "最高得分", "覆盖率"], cat_rows)

    # 将整体分析内容分段显示
    analysis_text = report["OverallAnalysis"]
    
    # 处理文本中的标记
    sections = analysis_text.split("【")
    for section in sections:
        if section.strip():
            if "】" in section:
                # 这是一个带标题的段落
                title, content = section.split("】", 1)
                # 添加小标题
                heading = doc.add_heading(title, level=2)
                # 添加内容
                paras = content.strip().split("\n")
                for para_text in paras:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.first_line_indent = Inches(0.5)
                        para.paragraph_format.space_after = Pt(8)
            else:
                # 普通段落
                paras = section.strip().split("\n")
                for para_text in paras:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.first_line_indent = Inches(0.5)
                        para.paragraph_format.space_after = Pt(8)
    
    # 添加分页符
    doc.add_page_break()

    # 各大类详细分析
    for idx, cat in enumerate(report["CategoryReports"]):
        doc.add_heading(cat["Category"], level=1)
        
        # 大类整体分析
        doc.add_heading("大类法规要求概述", level=2)
        para = doc.add_paragraph(cat["CategoryLawAnalysis"])
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.paragraph_format.space_after = Pt(12)
        
        # 子类别分析表格 - UPDATED SECTION
        doc.add_heading("子类别法规要求分析", level=2)
        
        # 创建表格 - 增加列数从4列到9列
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Light Grid Accent 1'
        table.autofit = False  # 关闭自动调整以便手动设置列宽
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        headers = [
            "子类别", 
            "覆盖情况", 
            "法规要求", 
            "关键条款",
            "所需制度文件",
            "所需管理体系",
            "禁止事项",
            "所需报告",
            "所需信息资料"
        ]
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # 设置表头格式
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置表头背景色
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E0E0E0')
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # 添加数据行
        for sub, info in cat["SubCategoryAnalysis"].items():
            row_cells = table.add_row().cells
            
            # 0. 子类别名称
            row_cells[0].text = sub
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 1. 覆盖情况 - 根据覆盖程度设置颜色
            coverage = info.get("Coverage", "未覆盖")
            row_cells[1].text = coverage
            para = row_cells[1].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.bold = True
            if coverage == "完全覆盖":
                run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif coverage == "部分覆盖":
                run.font.color.rgb = RGBColor(255, 140, 0)  # 橙色
            else:  # 未覆盖
                run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            
            # 2. 法规要求
            row_cells[2].text = info.get("LawRequirements", "")
            
            # 3. 关键条款
            key_provisions = info.get("KeyProvisions", [])
            if key_provisions:
                row_cells[3].text = "\n".join(f"• {p}" for p in key_provisions)
            else:
                row_cells[3].text = "无"
            
            # 4. 所需制度文件
            policy_docs = info.get("policy_document_needed", [])
            if policy_docs:
                row_cells[4].text = "\n".join(f"• {p}" for p in policy_docs)
            else:
                row_cells[4].text = "无"
            
            # 5. 所需管理体系
            mgmt_systems = info.get("management_system_needed", [])
            if mgmt_systems:
                row_cells[5].text = "\n".join(f"• {p}" for p in mgmt_systems)
            else:
                row_cells[5].text = "无"
            
            # 6. 禁止事项
            prohibited = info.get("Prohibited_List", [])
            if prohibited:
                row_cells[6].text = "\n".join(f"• {p}" for p in prohibited)
            else:
                row_cells[6].text = "无"
            
            # 7. 所需报告
            reports_needed = info.get("Report_needed", [])
            if reports_needed:
                row_cells[7].text = "\n".join(f"• {p}" for p in reports_needed)
            else:
                row_cells[7].text = "无"
            
            # 8. 所需信息资料
            data_info = info.get("data_info_needed", [])
            if data_info:
                row_cells[8].text = "\n".join(f"• {p}" for p in data_info)
            else:
                row_cells[8].text = "无"
            
            # 设置所有单元格的字体和对齐
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(9)
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.space_before = Pt(3)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # 设置列宽 - 调整为9列
        column_widths = [1.2, 0.8, 2.0, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2]  # 总宽度约10.2英寸
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
        
        # 为了更好的可读性，可以考虑将表格设置为横向页面
        # 或者将表格后的内容分成多个更小的表格
        
        # 可选：添加合规要点汇总（将原来在表格中的合规要点单独列出）
        if any(info.get("CompliancePoints") for info in cat["SubCategoryAnalysis"].values()):
            doc.add_heading("合规要点汇总", level=3)
            for sub, info in cat["SubCategoryAnalysis"].items():
                compliance_points = info.get("CompliancePoints", [])
                if compliance_points:
                    para = doc.add_paragraph()
                    para.add_run(f"{sub}：").bold = True
                    para.add_run("\n" + "\n".join(f"  • {p}" for p in compliance_points))
                    para.paragraph_format.space_after = Pt(6)
        
        # 大类合规指导
        doc.add_heading("合规措施建议", level=2)
        para = doc.add_paragraph(cat["CategoryComplianceGuidance"])
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.paragraph_format.space_after = Pt(12)
        
        # 在每个大类后添加一些空间，但不要分页（除非是最后一个）
        if idx < len(report["CategoryReports"]) - 1:
            doc.add_paragraph()
            doc.add_paragraph()

    # 添加报告信息
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # 设置所有表格的默认字体
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.font.name:
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存文档
    doc.save(out_file)


# ───────────────────────── 文本报告导出 ─────────────────────────
def _export_text_report(json_path: Path, out_file: Path):
    """基于综合分析结果生成定制化的文本报告"""
    # 读取JSON数据
    data = json.loads(json_path.read_text(encoding="utf-8"))
    reg_name = json_path.stem.replace("_综合分析结果", "")
    
    # 创建ID到名称的映射
    id_to_name, _ = _create_id_mappings()
    
    # 收集所有LLM的分数数据
    scores_by_llm = {}  # {llm_name: {req_id: score}}
    scores_by_category = collections.defaultdict(lambda: collections.defaultdict(list))  # {category: {llm: [scores]}}
    
    # 处理每个LLM的数据
    for llm_name, llm_data in data["LLM分析结果"].items():
        scores_by_llm[llm_name] = {}
        
        for cat_name, items in llm_data["详细分析"].items():
            for item in items:
                req_id = item.get("框架要求编号")
                if req_id and req_id in id_to_name:
                    # 计算分数（0-100）
                    coverage = item["法规覆盖情况"]
                    if coverage == "完全覆盖":
                        score = 100
                    elif coverage == "部分覆盖":
                        score = 75
                    elif coverage == "未覆盖":
                        score = 15
                    else:  # 未提及/不适用
                        score = 0
                    
                    scores_by_llm[llm_name][req_id] = score
                    scores_by_category[cat_name][llm_name].append(score)
    
    # 计算每个要求的平均分
    requirement_scores = []
    for req_id, info in id_to_name.items():
        scores = []
        for llm_name in scores_by_llm:
            if req_id in scores_by_llm[llm_name]:
                scores.append(scores_by_llm[llm_name][req_id])
            else:
                scores.append(0)  # 未提及的默认为0
        
        avg_score = sum(scores) / len(scores) if scores else 0
        requirement_scores.append({
            "id": req_id,
            "name": f"{req_id}. {info['name']}",
            "category": info['category'],
            "avg_score": avg_score,
            "scores": scores
        })
    
    # 按平均分降序排序
    requirement_scores.sort(key=lambda x: x["avg_score"], reverse=True)
    
    # 计算每个大类的统计数据
    category_stats = {}
    for cat_name in REGULATORY_FRAMEWORK.keys():
        all_scores = []
        for llm_name in scores_by_category[cat_name]:
            all_scores.extend(scores_by_category[cat_name][llm_name])
        
        if all_scores:
            category_stats[cat_name] = {
                "avg": np.mean(all_scores),
                "max": max(all_scores),
                "min": min(all_scores),
                "count": len([s for s in all_scores if s > 0])
            }
        else:
            category_stats[cat_name] = {
                "avg": 0,
                "max": 0,
                "min": 0,
                "count": 0
            }
    
    # 生成报告
    lines = []
    lines.append("=" * 80)
    lines.append(f"{reg_name} 法规合规覆盖分析报告")
    lines.append("=" * 80)
    
    # 一、类别分析
    lines.append("\n一、风险类别覆盖分析")
    lines.append("-" * 40)
    lines.append("类别名称                     平均分   最高分   最低分   覆盖项目数")
    lines.append("-" * 40)
    
    for cat_name, stats in category_stats.items():
        lines.append(f"{cat_name:<24} {stats['avg']:>6.1f}   {stats['max']:>6.0f}   {stats['min']:>6.0f}   {stats['count']:>10}")
    
    # 二、详细要求覆盖排名（完整列表）
    lines.append("\n\n二、详细要求覆盖排名（按平均分降序）")
    lines.append("-" * 80)
    lines.append("排名  要求编号及名称                                              平均分   覆盖情况")
    lines.append("-" * 80)
    
    for rank, req in enumerate(requirement_scores, 1):
        # 判断覆盖情况
        if req["avg_score"] >= 90:
            coverage = "完全覆盖"
        elif req["avg_score"] >= 60:
            coverage = "部分覆盖"
        elif req["avg_score"] >= 10:
            coverage = "低度覆盖"
        else:
            coverage = "未覆盖"
        
        lines.append(f"{rank:>3}   {req['name']:<50} {req['avg_score']:>6.1f}   {coverage}")
    
    # 三、重点发现
    lines.append("\n\n三、重点发现")
    lines.append("-" * 40)
    
    # 高覆盖要求
    high_coverage = [req for req in requirement_scores if req["avg_score"] >= 75]
    lines.append(f"\n高度覆盖的要求（≥75分）: {len(high_coverage)}项")
    for req in high_coverage[:5]:
        lines.append(f"  - {req['name']}: {req['avg_score']:.1f}分")
    
    # 中等覆盖要求
    medium_coverage = [req for req in requirement_scores if 25 <= req["avg_score"] < 75]
    lines.append(f"\n中度覆盖的要求（25-74分）: {len(medium_coverage)}项")
    
    # 低覆盖或未覆盖要求
    low_coverage = [req for req in requirement_scores if req["avg_score"] < 25]
    lines.append(f"\n低度覆盖或未覆盖的要求（<25分）: {len(low_coverage)}项")
    for req in low_coverage[:5]:
        lines.append(f"  - {req['name']}: {req['avg_score']:.1f}分")
    
    if len(low_coverage) > 5:
        lines.append(f"  ... 以及其他 {len(low_coverage) - 5} 项")
    
    # 四、合规建议重点
    lines.append("\n\n四、合规建议重点")
    lines.append("-" * 40)
    lines.append("基于覆盖分析，建议重点关注以下方面：")
    
    # 找出低分类别
    low_score_categories = [cat for cat, stats in category_stats.items() if stats['avg'] < 30]
    if low_score_categories:
        lines.append(f"\n1. 需要加强的风险类别：")
        for cat in low_score_categories:
            lines.append(f"   - {cat} (平均分: {category_stats[cat]['avg']:.1f})")
    
    # 找出最低分的5个要求
    lines.append(f"\n2. 急需建立或完善的制度（得分最低的5项）：")
    for req in requirement_scores[-5:]:
        lines.append(f"   - {req['name']} (得分: {req['avg_score']:.1f})")
    
    lines.append("\n" + "=" * 80)
    lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # 写入文件
    out_file.write_text("\n".join(lines), encoding="utf-8")


# ───────────────────────── 对外主函数 ─────────────────────────
def generate_overall_report(
    json_path: str | Path,
    model_cat="claude-opus-4-20250514",
    model_doc="claude-opus-4-20250514",
) -> tuple[Path, Path, Path]:
    """Returns (overall_json_path, overall_docx_path, analysis_txt_path)"""

    json_path = Path(json_path)
    cov, findings, advice, detailed_data = _gather(json_path)

    # 逐大类分析
    cat_reports = _build_category_reports(cov, findings, advice, detailed_data, model_cat)

    # 全文总体法规分析
    doc_prompt = textwrap.dedent(f"""
    你是一名资深法律分析专家，专注于企业境外投资合规法律分析。
    
    基于以下各大类的法规要求分析，请撰写一份综合分析报告。请直接输出纯文本内容，不要使用JSON格式。
    
    请按以下结构撰写（总计1000-1500字）：
    
    【法规整体分析】
    - 该法规的核心目的和规范重点
    - 法规对企业境外投资的主要要求领域
    - 法规的管理思路和监管重点
    - 法规要求建立的制度和管理体系
    - 法规禁止的负面清单
    - 法规要求或建议提交的报告
    - 法规要求提交的信息数据与资料
    
    【合规实施建议】
    - 基于法规要求，企业应建立的核心合规体系
    - 重点关注的合规领域和优先级
    - 实施这些合规要求的具体步骤建议
    
    各大类分析结果：
    {json.dumps(cat_reports, ensure_ascii=False)}
    
    请以专业、客观的语言撰写，重点说明法规要求了什么、企业应该做什么。
    直接输出文本内容，不要有任何JSON格式标记。
    """)
    
    overall_analysis = _call_anthropic(
        system_msg="你是专业的法律分析专家，请提供纯文本格式的专业分析，不要使用JSON格式。",
        user_msg=doc_prompt,
        model=model_doc,
        max_tokens=4000,
    )
    
    report = {
        "DocumentTitle": json_path.stem.replace("_综合分析结果", ""),
        "CategoryReports": cat_reports,
        "OverallAnalysis": overall_analysis.strip()
    }

    stem = json_path.stem
    out_json = json_path.parent / f"{stem}_overall.json"
    out_docx = json_path.parent / f"{stem}_overall.docx"
    reg_name = report["DocumentTitle"].replace('/', '_')
    out_txt = json_path.parent / f"{reg_name}_分析报告.txt"

    out_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    _export_word(report, out_docx, json_path.parent, json_path)
    _export_text_report(json_path, out_txt)

    return out_json, out_docx, out_txt


"""
quick_test_anthropic.py
测试 overall_reporter_anthropic.generate_overall_report()
"""

# 你的相对 JSON 路径
json_file = (
    "企业境外投资管理办法.json"
)

if __name__ == "__main__":
    rep_json, rep_docx, rep_txt = generate_overall_report(json_file)
    print("\n生成成功:")
    print("  •", rep_json)
    print("  •", rep_docx)
    print("  •", rep_txt)
